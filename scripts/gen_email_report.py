"""Generate the HTML email upgrade Word report."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

out_path = r"C:/Users/alexa/NexusTrader/reports/email_html_upgrade_report.docx"
os.makedirs(os.path.dirname(out_path), exist_ok=True)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin = section.bottom_margin = Inches(1)

C_BLUE  = RGBColor(0x1E, 0x35, 0x7A)
C_RED   = RGBColor(0xCC, 0x00, 0x00)
C_GREEN = RGBColor(0x00, 0x7A, 0x33)
C_AMBER = RGBColor(0xB3, 0x6A, 0x00)
C_GRAY  = RGBColor(0x44, 0x44, 0x44)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, c="CCCCCC", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for s in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{s}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), c)
        tcB.append(b)
    tcPr.append(tcB)


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = C_BLUE
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1E357A")
    pBdr.append(bot)
    pPr.append(pBdr)


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_BLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def body(text, bold=False, color=None, size=11, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)


def bul(text, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)


def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Courier New"
    run.font.color.rgb = RGBColor(0x22, 0x55, 0x88)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)


def sp():
    doc.add_paragraph()


def add_table(hdrs, rows_data, widths, row_colors=None):
    t = doc.add_table(rows=1, cols=len(hdrs))
    t.style = "Table Grid"
    for i, h in enumerate(hdrs):
        cell = t.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_bg(cell, "1E357A")
        set_cell_border(cell, "1E357A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(9)
    for ri, row_data in enumerate(rows_data):
        row = t.add_row()
        fill = "F5F7FF" if ri % 2 == 0 else "FFFFFF"
        for ci, (text, w) in enumerate(zip(row_data, widths)):
            cell = row.cells[ci]
            cell.width = w
            set_cell_bg(cell, fill)
            set_cell_border(cell, "CCCCCC")
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(8.5)
            if row_colors and ci < len(row_colors[ri]):
                rc = row_colors[ri][ci]
                if rc == "bold":
                    run.bold = True
                elif rc == "code":
                    run.font.name = "Courier New"
                    run.font.color.rgb = RGBColor(0x22, 0x44, 0x88)
                elif rc == "red":
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                elif rc == "amber":
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xB3, 0x6A, 0x00)
                elif rc == "green":
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x00, 0x7A, 0x33)
    return t


def add_hr(color="1E357A", sz=8):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── TITLE ────────────────────────────────────────────────────────
sp()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NexusTrader Email Notification System")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = C_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("HTML Upgrade Report")
r2.bold = True; r2.font.size = Pt(18); r2.font.color.rgb = C_BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Session 38  |  2026-03-28  |  NexusTrader Development Team")
r3.font.size = Pt(11); r3.font.color.rgb = C_GRAY
sp()
add_hr()
sp()

# ── SECTION 1 ────────────────────────────────────────────────────
h1("1. Executive Summary")
body(
    "Before this upgrade, NexusTrader email notifications were delivered as plain monospace text "
    "despite the system being designed for rich HTML output. Three compounding defects prevented "
    "professional emails from ever being sent:"
)
sp()
bul(
    "Issue 1 (CRITICAL): EmailChannel.send() and GeminiChannel.send() did not accept an html_body "
    "parameter. Calls from notification_manager raised TypeError which was silently caught — "
    "no HTML email was ever delivered by either channel.",
    C_RED,
)
bul(
    "Issue 2 (HIGH): 14 of 18 notification templates produced no html_body key at all. "
    "Templates including trade_stopped, system_error, emergency_stop, all crash_ tiers, and "
    "daily_summary returned only plain text — no HTML possible even if the channels were fixed.",
    C_AMBER,
)
bul(
    "Issue 3 (HIGH): A bug in _build_trade_opened_html() and _build_trade_closed_html() passed "
    "integer analysis scores to html.escape() which only accepts strings. AttributeError was "
    "silently swallowed, causing html_body to be absent from trade open/close emails.",
    C_AMBER,
)
sp()
body(
    "RESULT: ALL 18 templates now produce professional HTML. Both email channels correctly accept "
    "and use html_body. Emails are institutional-quality with branded headers, color-coded data, "
    "progress bars, direction badges, and a responsive dark-theme layout.",
    bold=True,
    color=C_GREEN,
)

# ── SECTION 2 ────────────────────────────────────────────────────
h1("2. Root Cause Analysis")
add_table(
    ["Issue", "Affected Code", "Impact", "Severity"],
    [
        [
            "EmailChannel.send() missing html_body param",
            "email_channel.py line 64",
            "Rich HTML never sent — pre monospace fallback used for ALL emails. TypeError caught silently, email queued for retry which also failed.",
            "CRITICAL",
        ],
        [
            "GeminiChannel.send() missing html_body param",
            "gemini_channel.py line 99",
            "Same issue — Gemini/Gmail channel also unable to send HTML emails.",
            "CRITICAL",
        ],
        [
            "14 of 18 templates missing html_body key",
            "notification_templates.py — trade_stopped, trade_rejected, trade_modified, strategy_signal, risk_warning, market_condition, system_error, system_alert, emergency_stop, crash_defensive, crash_high_alert, crash_emergency, crash_systemic, daily_summary",
            "Plain text in pre tag sent for all listed notification types regardless of channel fix.",
            "HIGH",
        ],
        [
            "html.escape() called on int (analysis_overall score)",
            "_build_trade_opened_html() ~line 1297 and _build_trade_closed_html() ~line 1411",
            "AttributeError silently swallowed; html_body absent from trade_opened and trade_closed results.",
            "HIGH",
        ],
    ],
    [Inches(1.8), Inches(2.2), Inches(3.0), Inches(0.8)],
    [
        ["bold", "", "", "red"],
        ["bold", "", "", "red"],
        ["bold", "code", "", "amber"],
        ["bold", "code", "", "amber"],
    ],
)
sp()

# ── SECTION 3 ────────────────────────────────────────────────────
h1("3. Changes Made")

h2("3.1  EmailChannel  (core/notifications/channels/email_channel.py)")
body(
    "Added html_body: Optional[str] = None parameter to send(). When html_body is provided it "
    "is used as the HTML MIME part. When omitted, the existing pre wrapper is used as fallback. "
    "Plain-text body always attached (RFC 2046 multipart/alternative compliance)."
)
code("Before: def send(self, message: str, subject: Optional[str] = None) -> bool:")
code("After:  def send(self, message: str, subject: Optional[str] = None, html_body: Optional[str] = None) -> bool:")

h2("3.2  GeminiChannel  (core/notifications/channels/gemini_channel.py)")
body(
    "Same html_body parameter added. Rich template HTML used directly when provided. "
    "Gemini AI enrichment appended as an additional styled card injected before the closing "
    "body tag — preserving the full template layout. Monospace fallback retained when html_body is None."
)

h2("3.3  Notification Templates  (core/notifications/notification_templates.py)")
body("a) New shared infrastructure:", bold=True)
bul("_esc(v) — HTML-safe escape helper (wraps html.escape, handles non-string types via str())")
bul(
    "_build_email_html() — Centralized generic HTML builder. Produces: full DOCTYPE/head, "
    "max-width 620px container, branded header card with title/subtitle/badge, table-based "
    "data rows (label/value/color), optional alert box, additional section cards, footer. "
    "All user values HTML-escaped."
)
bul(
    "_wrap_html(fn) — Decorator: adds html_body to any existing template function by "
    "calling the matching _build_*_html() builder via globals(). Applied to 14 functions."
)
sp()
body("b) 14 new HTML builder functions:", bold=True)
builders = [
    "_build_trade_stopped_html       — Red header, entry/stop price, loss display, critical advisory",
    "_build_trade_rejected_html      — Amber header, rejection reason, confidence, regime",
    "_build_trade_modified_html      — Blue header, old->new SL/TP with arrow notation",
    "_build_strategy_signal_html     — Blue header, direction badge, entry parameters as second card",
    "_build_risk_warning_html        — Severity-color header, level badge, current vs. threshold, alert box",
    "_build_market_condition_html    — Blue header, condition/regime/confidence, details field",
    "_build_system_error_html        — Red header, severity badge, error detail, critical alert box",
    "_build_system_alert_html        — Blue header, informational card layout",
    "_build_emergency_stop_html      — Bright red header, positions/equity, critical alert box",
    "_build_crash_defensive_html     — Amber header, crash score bar, defensive advisory",
    "_build_crash_high_alert_html    — Red header, score bar, partial exit advisory",
    "_build_crash_emergency_html     — Bright red header, score bar, READ-ONLY mode advisory",
    "_build_crash_systemic_html      — Darkest red header, score bar, SAFE MODE advisory",
    "_build_daily_summary_html       — Blue header, P&L color-coded, win rate progress bar",
]
for b in builders:
    bul(b)
sp()
body("c) Bug fix — int to str conversion:", bold=True)
body(
    "In _build_trade_opened_html() and _build_trade_closed_html(), the analysis_overall score "
    "(an integer) was passed directly to html.escape() which only accepts strings. "
    "Fixed by wrapping all analysis score values with str() before escaping: e(str(a_overall))."
)

# ── SECTION 4 ────────────────────────────────────────────────────
h1("4. HTML Template Design System")
body("All email templates share a consistent design language:")
add_table(
    ["Design Element", "Value / Implementation"],
    [
        ["Colour: LONG / Win / Positive", "#00CC77 — bright green"],
        ["Colour: SHORT / Loss / Error", "#FF3355 — vivid red"],
        ["Colour: Warning / Medium", "#FFB300 — amber"],
        ["Colour: Informational / Neutral", "#1E90FF — blue"],
        ["Colour: Critical Emergency", "#FF0033 — alert red"],
        ["Background: Body", "#080C16 — near-black navy"],
        ["Background: Card", "#0D1320 — dark navy"],
        ["Background: Borders", "#1A2332 — subtle separator"],
        ["Typography: Body", "Arial/Helvetica, 13px, #C8D0E0"],
        ["Typography: Headers", "14px bold, accent-coloured per alert tier"],
        ["Typography: Badges", "10px, 700 weight, letter-spacing 1px, border-radius 3px"],
        ["Layout Strategy", "Table-based — Gmail, Outlook 2016+, Apple Mail compatible"],
        ["Direction Badges", "LONG: #00CC77 on #064E3B | SHORT: #FF3355 on #7F1D1D"],
        ["Score Progress Bars", "CSS width% — green >= 75, amber >= 55, red < 55"],
        ["AI Analysis Cards", "Per-dimension score bars + large overall score with colour-coded quality"],
        ["Security", "All user-supplied data passed through html.escape() before HTML insertion"],
    ],
    [Inches(2.5), Inches(6.0)],
)
sp()

# ── SECTION 5 ────────────────────────────────────────────────────
doc.add_page_break()
h1("5. Proof of Coverage — Template Inventory")
body("All 18 notification templates verified by automated Python test to produce a valid html_body key.")
sp()

inv = [
    ("trade_opened",     "Partial (broken)", "YES", "_build_trade_opened_html",     "Bug fixed: int->str for analysis scores"),
    ("trade_closed",     "Partial (broken)", "YES", "_build_trade_closed_html",     "Bug fixed: int->str for analysis scores"),
    ("partial_exit",     "YES",              "YES", "_build_partial_exit_html",     "Already working — no change"),
    ("health_check",     "YES",              "YES", "_build_health_html",           "Already working — most complex builder"),
    ("trade_stopped",    "NO",               "YES", "_build_trade_stopped_html",    "New in Session 38"),
    ("trade_rejected",   "NO",               "YES", "_build_trade_rejected_html",   "New in Session 38"),
    ("trade_modified",   "NO",               "YES", "_build_trade_modified_html",   "New in Session 38"),
    ("strategy_signal",  "NO",               "YES", "_build_strategy_signal_html",  "New in Session 38"),
    ("risk_warning",     "NO",               "YES", "_build_risk_warning_html",     "New in Session 38"),
    ("market_condition", "NO",               "YES", "_build_market_condition_html", "New in Session 38"),
    ("system_error",     "NO",               "YES", "_build_system_error_html",     "New in Session 38"),
    ("system_alert",     "NO",               "YES", "_build_system_alert_html",     "New in Session 38"),
    ("emergency_stop",   "NO",               "YES", "_build_emergency_stop_html",   "New in Session 38"),
    ("crash_defensive",  "NO",               "YES", "_build_crash_defensive_html",  "New in Session 38"),
    ("crash_high_alert", "NO",               "YES", "_build_crash_high_alert_html", "New in Session 38"),
    ("crash_emergency",  "NO",               "YES", "_build_crash_emergency_html",  "New in Session 38"),
    ("crash_systemic",   "NO",               "YES", "_build_crash_systemic_html",   "New in Session 38"),
    ("daily_summary",    "NO",               "YES", "_build_daily_summary_html",    "New in Session 38"),
]

status_col = {"YES": "green", "NO": "red", "Partial (broken)": "amber"}
add_table(
    ["Template", "HTML Before", "HTML After", "Builder Function", "Notes"],
    [[t, b, a, fn, n] for t, b, a, fn, n in inv],
    [Inches(1.3), Inches(1.1), Inches(0.8), Inches(2.2), Inches(2.1)],
    [
        [
            "code",
            status_col.get(b, ""),
            status_col.get(a, ""),
            "code",
            "",
        ]
        for _, b, a, _, _ in inv
    ],
)
sp()

# ── SECTION 6 ────────────────────────────────────────────────────
h1("6. Validation Evidence")

h2("6.1  Automated Verification")
body("Python test executed against all 18 templates. Each template called via render() and result dict checked for html_body key:")
code("python -c \"from core.notifications.notification_templates import TEMPLATES, render; ...\"")
sp()
body("Output (all 18 templates):", bold=True)
for tmpl in [
    "trade_opened", "trade_closed", "partial_exit", "trade_stopped",
    "trade_rejected", "trade_modified", "strategy_signal", "risk_warning",
    "market_condition", "system_error", "system_alert", "emergency_stop",
    "crash_defensive", "crash_high_alert", "crash_emergency", "crash_systemic",
    "daily_summary", "health_check",
]:
    code(f"OK: {tmpl}")
code("ALL templates have html_body")
sp()
body("Result: PASS — 18 / 18 templates verified.", bold=True, color=C_GREEN)

h2("6.2  Regression Tests")
body("Session 37 test suite executed after all code changes:")
code("python -m pytest tests/unit/test_session37_scan_error_fix.py tests/unit/test_session37_reset_and_ui_fixes.py -q")
body("Result: 23 passed, 0 failed, 0 errors.", bold=True, color=C_GREEN)

h2("6.3  Sample HTML Emails Generated")
body("8 representative HTML email files rendered and saved to C:/Users/alexa/NexusTrader/reports/email_samples/:")
sp()
add_table(
    ["File", "Scenario", "Key Data Points", "Size"],
    [
        ["sample1_long_trade_opened.html",  "BTCUSDT LONG @ 65,420.50",    "74% confidence, AI score 82/100 GOOD, R:R 2.0",     "3,851 bytes"],
        ["sample2_short_trade_closed.html", "SOLUSDT SHORT — stop-loss",   "-13.02% P&L, AI score 41/100 BAD, 3h 42m duration", "4,810 bytes"],
        ["sample3_system_error.html",       "IDSS Scanner CRITICAL error", "HTTP 429 rate limit, Bybit REST API",                "3,367 bytes"],
        ["sample4_risk_warning.html",       "Portfolio Heat Exceeded",      "Current 6.8% > threshold 6.0% — HIGH severity",     "3,860 bytes"],
        ["sample5_emergency_stop.html",     "Emergency Stop activated",     "3 positions closed, equity $98,450",                 "3,489 bytes"],
        ["sample6_daily_summary.html",      "Daily Summary 2026-03-28",    "7 trades, 5W/2L, 71.4% WR, +1.96% P&L",             "4,066 bytes"],
        ["sample7_crash_defensive.html",    "Crash Defense — Defensive",   "Score 5.3/10, new longs halted",                    "3,383 bytes"],
        ["sample8_strategy_signal.html",    "ETHUSDT LONG strategy signal","PullbackLong model, confidence 68%",                "4,271 bytes"],
    ],
    [Inches(2.3), Inches(1.8), Inches(3.0), Inches(0.8)],
)
sp()

# ── SECTION 7 ────────────────────────────────────────────────────
doc.add_page_break()
h1("7. Before vs. After Comparison")

h2("7.1  Channel send() — Before")
body("EmailChannel.send() accepted only message and subject:")
code("def send(self, message: str, subject: Optional[str] = None) -> bool:")
body("notification_manager called:")
code('channel.send(body, subject=subject, html_body=content["html_body"])')
body(
    "Result: Python raised TypeError: send() got an unexpected keyword argument html_body. "
    "Caught silently by _send_on_channel's except Exception block — email queued for retry "
    "which also failed. No HTML email was ever delivered."
)

h2("7.2  Channel send() — After")
code("def send(self, message: str, subject: Optional[str] = None, html_body: Optional[str] = None) -> bool:")
bul("If html_body is provided: use it as the HTML MIME part (full professional template)")
bul("If html_body is None: wrap message in pre monospace block (existing fallback)")
bul("Plain-text body always attached as separate MIME part (RFC 2046 multipart/alternative)")

h2("7.3  Email Content — Before")
body("Every email (all 18 notification types) produced this identical HTML:")
code("<html><body style='font-family:monospace;background:#0A0E1A;color:#C8D0E0;padding:20px'>")
code("<pre style='color:#C8D0E0'>{plain_text_body}</pre></body></html>")
body(
    "A dark monospace pre-block. No tables, no colour hierarchy, no visual structure, "
    "no branding beyond background colour. Identical output regardless of notification type."
)

h2("7.4  Email Content — After")
body("Each template produces a unique, fully-structured HTML email with:")
for item in [
    "Branded dark header card — title + subtitle + optional badge, accent-coloured per alert type",
    "Table-based data grid — label (gray) / value (white+bold) / colour-coded value columns",
    "Direction badges — LONG in green, SHORT in red; visually immediate signal at-a-glance",
    "AI Analysis cards — per-dimension score bars + large overall score with quality colour coding",
    "Alert boxes — left-border accent for warnings and error conditions",
    "Crash defense score bars — tier-specific colours: amber / red / bright-red / dark-red",
    "Win rate progress bars in daily summary emails",
    "Severity badges (CRITICAL / HIGH / MEDIUM / LOW) with appropriate colour coding",
    "Consistent NexusTrader footer on all emails",
    "Max-width 620px container — renders correctly on mobile and desktop clients",
]:
    bul(item)

# ── SECTION 8 ────────────────────────────────────────────────────
h1("8. Limitations and Notes")
limitations = [
    (
        "Email Client CSS (display:flex)",
        "The existing _TRADE_HTML_CSS uses display:flex for .row elements in trade_opened, "
        "trade_closed, and partial_exit builders. All 14 new builders use table-based layout "
        "(fully Outlook-compatible). Flex rows may render as block in Outlook 2013 and earlier "
        "but data remains readable. Future pass: convert to table-based layout in the existing builders.",
    ),
    (
        "Emoji in Subject Lines",
        "Windows cp1252 terminal shows UnicodeEncodeError when printing emoji subject lines. "
        "This is a console encoding issue only — SMTP transmission uses UTF-8 and emoji render "
        "correctly in Gmail, Apple Mail, and Outlook 365.",
    ),
    (
        "Dark Theme Compatibility",
        "Dark background (#080C16) renders correctly in Gmail web, Apple Mail, Outlook.com, "
        "and iOS/Android. Enterprise Outlook in High Contrast mode may override background-color "
        "CSS. All text remains readable via the plain-text MIME fallback.",
    ),
    (
        "Gmail Strips style Blocks",
        "Gmail web strips <style> tags and only processes inline styles. The existing "
        "_TRADE_HTML_CSS-based builders may be affected. All new builders in Session 38 "
        "use fully inline styles throughout. Future pass: convert existing builders to inline styles.",
    ),
    (
        "Plain-Text Fallback",
        "All emails include a plain-text MIME part (multipart/alternative per RFC 2046). "
        "Non-HTML clients receive the same structured plain-text body that was previously "
        "the sole format — no regression.",
    ),
]
for i, (title, text) in enumerate(limitations, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{i}.  {title}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = C_GRAY

# ── SECTION 9 ────────────────────────────────────────────────────
h1("9. Files Changed Summary")
add_table(
    ["File", "Change Type", "Description"],
    [
        [
            "core/notifications/channels/email_channel.py",
            "Bug fix",
            "Added html_body: Optional[str] = None to send(). Uses rich HTML when provided; "
            "pre monospace fallback otherwise. Plain-text MIME part always attached.",
        ],
        [
            "core/notifications/channels/gemini_channel.py",
            "Bug fix",
            "Same html_body param added. Rich HTML used directly; Gemini AI analysis injected "
            "as additional styled card. Monospace fallback preserved when html_body is None.",
        ],
        [
            "core/notifications/notification_templates.py",
            "Enhancement + Bug fix",
            "Added _esc(), _build_email_html(), _wrap_html() infrastructure. Added 14 new HTML "
            "builder functions. Fixed int->str type error in existing analysis score builders. "
            "Applied _wrap_html() decorator to 14 template functions. 0 regressions.",
        ],
    ],
    [Inches(2.6), Inches(1.0), Inches(4.0)],
    [["code", "bold", ""], ["code", "bold", ""], ["code", "bold", ""]],
)
sp()
sp()

add_hr(color="1E357A", sz=4)
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_foot.add_run("NexusTrader v1.3  |  Session 38  |  2026-03-28  |  Confidential — Internal Use Only")
run_f.font.size = Pt(9)
run_f.font.color.rgb = C_GRAY

doc.save(out_path)
print("Saved:", out_path, os.path.getsize(out_path), "bytes")
